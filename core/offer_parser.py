"""
InternShield — Offer Letter Parser
=====================================
Extracts text from PDF / DOCX offer letters and runs deep forensic analysis.
"""

import re
import os
from typing import Dict, Any, List
from datetime import datetime

from . import config


class OfferLetterParser:

    def parse(self, filepath: str) -> Dict[str, Any]:
        """Parse and analyze an offer letter file (PDF or DOCX)."""
        text = ""
        metadata = {}
        
        if filepath.lower().endswith(".pdf"):
            text = self._read_pdf(filepath)
            metadata = self._read_pdf_metadata(filepath)
        elif filepath.lower().endswith((".docx", ".doc")):
            text = self._read_docx(filepath)
            metadata = self._read_docx_properties(filepath)
        else:
            return {"error": "Unsupported file format"}

        if not text:
            return {"error": "Could not extract text from file"}

        return self._analyze(text, filepath, metadata)

    # ──────────────────────────────────────────────────────────────────────────
    # READERS & FORENSICS
    # ──────────────────────────────────────────────────────────────────────────
    @staticmethod
    def _read_pdf(path: str) -> str:
        try:
            import PyPDF2
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                pages  = [p.extract_text() or "" for p in reader.pages]
                return "\n".join(pages)
        except Exception: return ""

    @staticmethod
    def _read_pdf_metadata(path: str) -> Dict:
        try:
            import PyPDF2
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                meta = reader.metadata
                return {
                    "author":    str(meta.get('/Author', 'None')),
                    "creator":   str(meta.get('/Creator', 'None')),
                    "producer":  str(meta.get('/Producer', 'None')),
                    "created":   str(meta.get('/CreationDate', 'None')),
                }
        except Exception: return {}

    @staticmethod
    def _read_docx(path: str) -> str:
        try:
            from docx import Document
            doc   = Document(path)
            paras = [p.text for p in doc.paragraphs]
            return "\n".join(paras)
        except Exception: return ""

    @staticmethod
    def _read_docx_properties(path: str) -> Dict:
        try:
            from docx import Document
            doc   = Document(path)
            prop  = doc.core_properties
            return {
                "author":   str(prop.author or 'None'),
                "creator":  str(prop.last_modified_by or 'None'),
                "producer": "Microsoft Word (DOCX)",
                "created":  str(prop.created or 'None'),
            }
        except Exception: return {}

    # ──────────────────────────────────────────────────────────────────────────
    # ANALYSIS ENGINE
    # ──────────────────────────────────────────────────────────────────────────
    def _analyze(self, text: str, filepath: str, metadata: Dict) -> Dict[str, Any]:
        lower = text.lower()
        result = {
            "filepath":           filepath,
            "filename":           os.path.basename(filepath),
            "full_text":          text,
            "metadata":           metadata,
            "company_name":       self._extract_company(text),
            "salary":             self._extract_salary(text),
            "role":               self._extract_role(text),
            "joining_date":       self._extract_joining_date(text),
            "contact_info":       self._extract_contacts(text),
            "pressure_tactics":   self._find_pressure(lower),
            "financial_flags":    self._find_financial_flags(lower),
            "fake_signals":       self._find_fake_signals(lower),
            "suspicious_clauses": self._find_suspicious_clauses(lower),
            "grammar_issues":     self._detect_grammar_issues(text),
            "risk_signals":       [],
            "offer_score":        0,
        }

        # 🔍 METADATA AUDIT
        author   = metadata.get("author", "").lower()
        producer = metadata.get("producer", "").lower()
        company  = result["company_name"].lower()

        if company and author != "none" and company not in author:
            if any(x in author for x in ["admin", "user", "owner", "desktop"]):
                pass # common generic authors
            else:
               result["risk_signals"].append(f"Metadata Mismatch: Document author is '{author}' but company is '{result['company_name']}'")

        if any(x in producer for x in ["wps", "libre", "openoffice"]):
            if company and any(y in company for y in ["google", "microsoft", "amazon", "apple"]):
                result["risk_signals"].append(f"Infrastructure Mismatch: Document created in '{metadata['producer']}' but claims to be from '{result['company_name']}' (Red Flag)")

        # Aggregate risk signals
        if result["pressure_tactics"]:
            result["risk_signals"].append(f"Pressure tactics found: {result['pressure_tactics']}")
        if result["financial_flags"]:
            result["risk_signals"].append(f"⚠️ FINANCIAL RED FLAGS — they may ask you to PAY: {result['financial_flags']}")
        if result["fake_signals"]:
            result["risk_signals"].append(f"Fake offer signals: {result['fake_signals']}")
        if result["suspicious_clauses"]:
            result["risk_signals"].append(f"Suspicious clauses: {result['suspicious_clauses']}")
        if result["grammar_issues"] > 5:
            result["risk_signals"].append(f"High grammar/spelling density — likely template spam")
        if not result["company_name"]:
            result["risk_signals"].append("No clear company name found in offer letter")

        result["offer_score"] = self._compute_offer_score(result)
        return result

    # ──────────────────────────────────────────────────────────────────────────
    # EXTRACTORS & HELPERS
    # ──────────────────────────────────────────────────────────────────────────
    @staticmethod
    def _extract_company(text: str) -> str:
        patterns = [
            r"(?:from|by|company name[:\s]+|organization[:\s]+)([A-Z][A-Za-z\s&.,]+(?:Pvt\.?\s*Ltd\.?|LLC|Inc\.?|Ltd\.?|LLP|Technologies|Solutions|Services|Systems|Consulting|Group|Corp\.?))",
            r"([A-Z][A-Za-z\s&.,]+(?:Pvt\.?\s*Ltd\.?|LLC|Inc\.?|Ltd\.?|LLP|Technologies|Solutions|Services|Systems|Consulting|Group|Corp\.?))",
        ]
        for pat in patterns:
            m = re.search(pat, text)
            if m: return m.group(1).strip()
        return ""

    @staticmethod
    def _extract_salary(text: str) -> str:
        patterns = [
            r"(?:stipend|salary|ctc|compensation)[:\s]+(?:INR|Rs\.?|₹)?\s*([\d,]+(?:\.\d+)?)\s*(?:per month|/month|p\.?m\.?|lpa|per annum)?",
            r"(?:INR|Rs\.?|₹)\s*([\d,]+(?:\.\d+)?)\s*(?:per month|/month|p\.?m\.?|lpa|per annum)?",
        ]
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m: return m.group(1).replace(",", "")
        return ""

    @staticmethod
    def _extract_role(text: str) -> str:
        patterns = [r"(?:position|role|designation|post)[:\s]+([A-Za-z\s/-]+)(?:\n|,|\.|;)"]
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m: return m.group(1).strip()
        return ""

    @staticmethod
    def _extract_joining_date(text: str) -> str:
        patterns = [r"(?:joining date|join by|start date|commence)[:\s]+([A-Za-z0-9,\s./-]+)"]
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m: return m.group(1).strip()
        return ""

    @staticmethod
    def _extract_contacts(text: str) -> Dict:
        emails = re.findall(r'[\w.+-]+@[\w.-]+\.\w+', text)
        phones = re.findall(r'(?:\+91[\s-]?)?\d{10}', text)
        return {"emails": emails, "phones": list(set(phones))}

    def _find_pressure(self, lower: str) -> List[str]:
        found = []
        for phrase in config.PRESSURE_TACTICS:
            if phrase in lower: found.append(phrase)
        return found

    def _find_financial_flags(self, lower: str) -> List[str]:
        found = []
        for phrase in config.FINANCIAL_RED_FLAGS:
            if phrase in lower: found.append(phrase)
        return found

    def _find_fake_signals(self, lower: str) -> List[str]:
        found = []
        for phrase in config.FAKE_OFFER_SIGNALS:
            if phrase in lower: found.append(phrase)
        return found

    @staticmethod
    def _find_suspicious_clauses(lower: str) -> List[str]:
        clauses = []; suspicious = ["share your aadhar", "share your pan", "training bond", "security bond", "penalty for early exit"]
        for c in suspicious:
            if c in lower: clauses.append(c)
        return clauses

    @staticmethod
    def _detect_grammar_issues(text: str) -> int:
        issues = 0
        issues += len(re.findall(r'  +', text))
        issues += len(re.findall(r'[!?]{2,}', text))
        return issues

    def _compute_offer_score(self, result: Dict) -> int:
        score = 0
        score += len(result["pressure_tactics"])  * 20
        score += len(result["financial_flags"])   * 40
        score += len(result["fake_signals"])       * 15
        score += len(result["suspicious_clauses"]) * 25
        # Metadata penalties
        if any("Metadata Mismatch" in s for s in result["risk_signals"]): score += 20
        if any("Infrastructure Mismatch" in s for s in result["risk_signals"]): score += 30
        return min(score, 100)
